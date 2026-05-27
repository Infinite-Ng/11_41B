"""
11.41B Status Report Generator
Flask web app that reads MDB and history_space files and generates
a Word document report from the standard template.
"""
import os
import io
import re
import copy
import glob
from datetime import datetime

import pyodbc
import openpyxl
from flask import Flask, render_template, request, send_file, jsonify
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

app = Flask(__name__)

# ── Paths ─────────────────────────────────────────────────────────────────────
BASE_DIR     = r'M:\BRSSD\CSS\11.41B\2026'
TEMPLATE_PATH = r'M:\BRSSD\CSS\11.41B\11.41B_WM1668_D72317_Draft - sample.docx'

# ── Helpers ───────────────────────────────────────────────────────────────────

def ordinal_suffix(n: int) -> str:
    """Return the ordinal suffix for an integer (e.g. 1→'st', 2→'nd', ...)."""
    if 11 <= (n % 100) <= 13:
        return 'th'
    return {1: 'st', 2: 'nd', 3: 'rd'}.get(n % 10, 'th')


def ordinal(n: int) -> str:
    return f'{n}{ordinal_suffix(n)}'


def _numeric_key(v: str):
    """Sort provision values numerically: '9.7' < '9.13' < '9.14'."""
    try:
        parts = v.split('.')
        return tuple(int(p) for p in parts)
    except ValueError:
        return (0, v)


def _make_run_xml(
    text: str,
    bold: bool = False,
    highlight: str = None,
    superscript: bool = False,
    font_name: str = 'Calibri',
    sz_cs: int = 22,
) -> object:
    """
    Create a ``w:r`` OxmlElement.
    Matches the body-text run style of the template:
      rFonts(Calibri) + optional b/bCs + optional vertAlign(superscript)
      + optional highlight + szCs(22).
    """
    r   = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    # Font family (ascii + hAnsi + cs)
    if font_name:
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:ascii'), font_name)
        rFonts.set(qn('w:hAnsi'), font_name)
        rFonts.set(qn('w:cs'),    font_name)
        rPr.append(rFonts)

    # Bold
    if bold:
        rPr.append(OxmlElement('w:b'))
        rPr.append(OxmlElement('w:bCs'))

    # Superscript
    if superscript:
        va = OxmlElement('w:vertAlign')
        va.set(qn('w:val'), 'superscript')
        rPr.append(va)

    # Highlight
    if highlight:
        hl = OxmlElement('w:highlight')
        hl.set(qn('w:val'), highlight)
        rPr.append(hl)

    # Complex-script size (matches template's szCs val="22")
    if sz_cs:
        szCs = OxmlElement('w:szCs')
        szCs.set(qn('w:val'), str(sz_cs))
        rPr.append(szCs)

    r.append(rPr)

    # w:t
    t = OxmlElement('w:t')
    t.text = text
    if text != text.strip() or not text:
        t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    r.append(t)
    return r


# ── Data readers ──────────────────────────────────────────────────────────────

def get_mdb_data(notice_id: str) -> dict:
    """
    Open the MDB at BASE_DIR/<notice_id>/<notice_id>.mdb (read-only) and
    return:
      com_el: list of dicts with ntc_id, adm, sat_name, type, d_rcv
      still_exist: 'YES' if provn has any non-O agree_st, else 'NO'
    """
    mdb_path = os.path.join(BASE_DIR, notice_id, f'{notice_id}.mdb')
    if not os.path.exists(mdb_path):
        raise FileNotFoundError(f'MDB file not found: {mdb_path}')

    conn_str = (
        f'Driver={{Microsoft Access Driver (*.mdb, *.accdb)}};'
        f'DBQ={mdb_path};ReadOnly=True;'
    )
    conn = pyodbc.connect(conn_str)
    cursor = conn.cursor()

    # com_el ─────────────────────────────────────────────────────────────────
    cursor.execute('SELECT ntc_id, adm, sat_name, long_nom, d_rcv FROM com_el')
    com_el_rows = []
    for row in cursor.fetchall():
        d_rcv = row[4].strftime('%d.%m.%Y') if row[4] else ''
        com_el_rows.append({
            'ntc_id':   str(row[0]) if row[0] is not None else '',
            'adm':      (row[1] or '').strip(),
            'sat_name': (row[2] or '').strip(),
            'type':     'G' if row[3] is not None else 'N',
            'd_rcv':    d_rcv,
        })

    # provn.coord_prov: YES if any entry contains '11.41' ────────────────────
    cursor.execute("SELECT COUNT(*) FROM provn WHERE coord_prov LIKE '%11.41%'")
    has_1141 = cursor.fetchone()[0]
    still_exist = 'YES' if has_1141 > 0 else 'NO'

    conn.close()
    return {'com_el': com_el_rows, 'still_exist': still_exist}


def get_history_data(notice_id: str) -> dict:
    """
    Find the **(history_space).xlsx** file under
    BASE_DIR/BR_TEXT_RESULTS/<notice_id>/Status Review/
    and return:
      provn_targets: sorted list of distinct provn_target strings
      adms:          sorted list of distinct adm strings
    """
    status_dir = os.path.join(
        BASE_DIR, 'BR_TEXT_RESULTS', notice_id, 'Status Review'
    )
    if not os.path.isdir(status_dir):
        raise FileNotFoundError(
            f'Status Review folder not found: {status_dir}'
        )

    # Find the file (may be .xlsx or .xls)
    matches = glob.glob(os.path.join(status_dir, '*(history_space).*'))
    if not matches:
        raise FileNotFoundError(
            f'No (history_space) file found in: {status_dir}'
        )
    xlsx_path = matches[0]

    wb = openpyxl.load_workbook(xlsx_path, read_only=True)
    ws = wb.active

    # Row 1 = sheet title  (e.g. 'COMPASS-80E')
    # Row 2 = column headers
    # Row 3+ = data
    header_row = [
        (str(h).strip() if h is not None else '')
        for h in next(ws.iter_rows(min_row=2, max_row=2, values_only=True))
    ]

    def col_idx(name: str, default: int) -> int:
        try:
            return header_row.index(name)
        except ValueError:
            return default

    provn_target_idx = col_idx('provn_target', 8)
    adm_idx          = col_idx('adm', 9)

    provn_targets: set = set()
    adms:          set = set()

    for row in ws.iter_rows(min_row=3, values_only=True):
        if row[0] is None:
            break
        pt  = str(row[provn_target_idx]).strip() if row[provn_target_idx] is not None else ''
        adm = str(row[adm_idx]).strip()          if row[adm_idx]          is not None else ''
        if pt:
            provn_targets.add(pt)
        if adm:
            adms.add(adm)

    wb.close()

    return {
        'provn_targets': sorted(provn_targets, key=_numeric_key),
        'adms':          sorted(adms),
    }

def get_capture_data(notice_id: str) -> dict:
    """
    Find *(capture_space).xlsx in the Status Review folder and extract:
      findings_review:    'YES' if any IsReviewed=Yes row has a non-ANN 13A value;
                          'NO' if all IsReviewed=Yes rows have 13A=ANN (or no Yes rows)
      still_exist_suffix: letter (A/B) + digit (1/2/3)
        A = ALL rows are IsReviewed=Yes; B = at least one is not Yes
        1 = no Yes-row has 13A=ANN; 2 = mixed; 3 = all Yes-rows have 13A=ANN
    Returns empty strings for both fields if the file is not found.
    """
    status_dir = os.path.join(BASE_DIR, 'BR_TEXT_RESULTS', notice_id, 'Status Review')
    matches    = glob.glob(os.path.join(status_dir, '*(capture_space).*'))
    if not matches:
        return {'findings_review': '', 'still_exist_suffix': ''}

    wb = openpyxl.load_workbook(matches[0], read_only=True)
    ws = wb.active

    headers = [
        (str(h).strip() if h is not None else '')
        for h in next(ws.iter_rows(min_row=1, max_row=1, values_only=True))
    ]

    def find_col(*candidates):
        for name in candidates:
            try:
                return headers.index(name)
            except ValueError:
                pass
        return -1

    reviewed_idx = find_col('IsReviewed', 'Is Reviewed')
    col13a_idx   = find_col('13 A', '13A', '13a', '13 a')

    if reviewed_idx == -1 or col13a_idx == -1:
        wb.close()
        return {'findings_review': '', 'still_exist_suffix': ''}

    all_reviewed      = True
    reviewed_13a_vals = []

    for row in ws.iter_rows(min_row=2, values_only=True):
        if row[0] is None:
            break
        is_rev = (str(row[reviewed_idx]).strip() if row[reviewed_idx] is not None else '').lower()
        val13a = (str(row[col13a_idx]).strip()   if row[col13a_idx]   is not None else '')
        if is_rev != 'yes':
            all_reviewed = False
        else:
            reviewed_13a_vals.append(val13a)

    wb.close()

    letter = 'A' if all_reviewed else 'B'

    if not reviewed_13a_vals:
        number = ''
    else:
        ann_count     = sum(1 for v in reviewed_13a_vals if v.upper() == 'ANN')
        non_ann_count = len(reviewed_13a_vals) - ann_count
        if ann_count == 0:
            number = '1'      # no Yes-row has 13A=ANN
        elif non_ann_count == 0:
            number = '3'      # all Yes-rows have 13A=ANN
        else:
            number = '2'      # mixed

    findings_review = (
        'YES' if any(v.upper() != 'ANN' for v in reviewed_13a_vals) else 'NO'
    )
    return {
        'findings_review':    findings_review,
        'still_exist_suffix': f'{letter}{number}',
    }

# ── Document builders ─────────────────────────────────────────────────────────

def _update_meeting_para(para, wm_num: str, meeting_date: str) -> None:
    """
    Replace the 'Presented at the 1668th meeting on 4 June 2026' paragraph.
    Rebuilds as three runs so the ordinal suffix (th/st/nd/rd) is superscript:
      Run 1: 'Presented at the {wm_num}'  – bold, Calibri
      Run 2: '{suffix}'                   – bold, Calibri, superscript
      Run 3: ' meeting on {date}'         – bold, Calibri
    """
    W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    p_elem = para._p

    # Remove all existing w:r elements
    for r in list(p_elem.findall(f'{{{W}}}r')):
        p_elem.remove(r)

    # Determine the ordinal suffix
    try:
        wm_int = int(wm_num)
        suffix = ordinal_suffix(wm_int)
        num_str = str(wm_int)
    except (ValueError, TypeError):
        suffix  = 'th'
        num_str = str(wm_num) if wm_num else 'xxxx'

    # Run 1 – 'Presented at the <number>' (bold)
    p_elem.append(_make_run_xml(
        f'Presented at the {num_str}', bold=True
    ))
    # Run 2 – superscript ordinal suffix (bold + superscript)
    p_elem.append(_make_run_xml(
        suffix, bold=True, superscript=True
    ))
    # Run 3 – rest of the heading (bold)
    p_elem.append(_make_run_xml(
        f' meeting on {meeting_date}', bold=True
    ))


def _replace_provn_section(para, provn_targets: list) -> None:
    """
    In the 'obligation to coordinate' paragraph, find the section between
    '|X| to ' and 'or removed from' and replace it with highlighted runs
    for each distinct provn_target (format: 'No. <value> |O|').
    """
    W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    runs = para.runs

    # Find the run containing '|X| to' and the run containing 'or removed from'
    start_after = None   # index of the '|X| to ' run
    end_before  = None   # index of the 'or removed from' run

    for i, run in enumerate(runs):
        txt = run.text
        if '|X| to' in txt:
            start_after = i
        elif start_after is not None and 'or removed from' in txt:
            end_before = i
            break

    if start_after is None or end_before is None:
        return   # pattern not found; leave paragraph untouched

    # Collect the w:r XML elements that need to be deleted
    runs_to_delete = [runs[i]._r for i in range(start_after + 1, end_before)]

    p_elem   = para._p
    ref_elem = runs[start_after]._r

    # Delete old provn runs
    for r_elem in runs_to_delete:
        p_elem.remove(r_elem)

    # Insert new highlighted runs right after ref_elem
    ref_idx = list(p_elem).index(ref_elem)
    insert_pos = ref_idx + 1

    new_elems = []
    for j, target in enumerate(provn_targets):
        if j > 0:
            # separator between items
            new_elems.append(_make_run_xml(', ', highlight='yellow'))
        new_elems.append(_make_run_xml('No. ',    bold=False, highlight='yellow'))
        new_elems.append(_make_run_xml(target,    bold=True,  highlight='yellow'))
        new_elems.append(_make_run_xml(' |O|',    bold=False, highlight='yellow'))

    for offset, elem in enumerate(new_elems):
        p_elem.insert(insert_pos + offset, elem)


def _update_footer(doc, wm_num: str, d_num: str) -> None:
    """
    Update the 'D <number>' document reference in all section footers.
    Replaces 'D<optional-space><digits>' (not preceded by a letter) with
    'D {d_num}'. Other footer content (No.11.41B, page number) is unchanged.
    """
    for section in doc.sections:
        for para in section.footer.paragraphs:
            for run in para.runs:
                t = run.text
                if d_num:
                    t = re.sub(r'(?<![A-Za-z])D\s*\d+', f'D {d_num}', t)
                run.text = t


def _update_signoff_date(doc, date_str: str) -> None:
    """
    Find the last date-pattern paragraph in the document body and replace
    its date with date_str (e.g. '27 May 2026').
    Searches paragraphs from the end to catch the sign-off date.
    """
    date_re = re.compile(
        r'\b\d{1,2}\s+(?:January|February|March|April|May|June|July|'
        r'August|September|October|November|December)\s+\d{4}\b',
        re.IGNORECASE,
    )
    for para in reversed(doc.paragraphs):
        if date_re.search(para.text):
            for run in para.runs:
                if date_re.search(run.text):
                    run.text = date_re.sub(date_str, run.text)
            break


def _set_cell_text(cell, text: str, force_calibri: bool = False) -> None:
    """
    Clear a table cell and set plain text.
    Preserves the run properties (rPr) from the cloned template row so that
    font and size are maintained.
    Pass force_calibri=True to override rFonts with explicit Calibri
    (used for the 'Provision updated from 11.41|X|' column).
    """
    para   = cell.paragraphs[0]
    W      = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    p_elem = para._p

    existing_runs = p_elem.findall(f'{{{W}}}r')

    # Capture run properties from the first existing run
    existing_rPr = None
    if existing_runs:
        rPr_list = existing_runs[0].findall(f'{{{W}}}rPr')
        if rPr_list:
            existing_rPr = copy.deepcopy(rPr_list[0])

    # Remove all existing runs
    for r in existing_runs:
        p_elem.remove(r)

    # Build a new run
    r = OxmlElement('w:r')
    if existing_rPr is not None:
        rPr_to_use = existing_rPr
        if force_calibri:
            # Replace any existing rFonts with explicit Calibri
            for rf in list(rPr_to_use.findall(f'{{{W}}}rFonts')):
                rPr_to_use.remove(rf)
            rFonts = OxmlElement('w:rFonts')
            rFonts.set(qn('w:ascii'), 'Calibri')
            rFonts.set(qn('w:hAnsi'), 'Calibri')
            rFonts.set(qn('w:cs'),    'Calibri')
            rPr_to_use.insert(0, rFonts)
        r.append(rPr_to_use)
    else:
        # Fallback: always explicit Calibri
        rPr    = OxmlElement('w:rPr')
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:ascii'), 'Calibri')
        rFonts.set(qn('w:hAnsi'), 'Calibri')
        rFonts.set(qn('w:cs'),    'Calibri')
        rPr.append(rFonts)
        r.append(rPr)

    t = OxmlElement('w:t')
    t.text = text
    if text != text.strip():
        t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    r.append(t)
    p_elem.append(r)


def _set_cell_still_exist(cell, main_text: str, suffix: str) -> None:
    """
    Set a table cell to main_text with a superscript suffix.
    E.g. main_text='YES', suffix='A3' produces 'YES' + superscript 'A3'.
    Both runs preserve the template row's rPr (font + size).
    """
    para   = cell.paragraphs[0]
    W      = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    p_elem = para._p

    existing_runs = p_elem.findall(f'{{{W}}}r')
    base_rPr = None
    if existing_runs:
        rPr_list = existing_runs[0].findall(f'{{{W}}}rPr')
        if rPr_list:
            base_rPr = copy.deepcopy(rPr_list[0])
    for r in existing_runs:
        p_elem.remove(r)

    # Run 1: main text
    r1 = OxmlElement('w:r')
    if base_rPr is not None:
        r1.append(copy.deepcopy(base_rPr))
    t1 = OxmlElement('w:t')
    t1.text = main_text
    r1.append(t1)
    p_elem.append(r1)

    # Run 2: superscript suffix
    if suffix:
        r2   = OxmlElement('w:r')
        rPr2 = copy.deepcopy(base_rPr) if base_rPr is not None else OxmlElement('w:rPr')
        va   = OxmlElement('w:vertAlign')
        va.set(qn('w:val'), 'superscript')
        rPr2.append(va)
        r2.append(rPr2)
        t2 = OxmlElement('w:t')
        t2.text = suffix
        r2.append(t2)
        p_elem.append(r2)


def _replace_table_data(table, notice_data_list: list) -> None:
    """
    Remove all sample data rows (row index >= 1) and add one row per com_el
    record across all notices, cloning the first data row's XML for formatting.
    notice_data_list: list of (mdb_data, history_data, capture_data) tuples.
    """
    # Save the first data row as formatting template (before deletion)
    template_tr = None
    if len(table.rows) > 1:
        template_tr = copy.deepcopy(table.rows[1]._tr)

    # Delete all data rows (keep header row 0)
    while len(table.rows) > 1:
        table._tbl.remove(table.rows[1]._tr)

    for mdb_data, history_data, capture_data in notice_data_list:
        adm_coord          = ', '.join(history_data['adms'])
        provn_col          = ', '.join(
            f'No. {pt} |O|' for pt in history_data['provn_targets']
        )
        findings_review    = capture_data.get('findings_review', '')
        still_exist_suffix = capture_data.get('still_exist_suffix', '')

        for record in mdb_data['com_el']:
            # Clone template row or add a new one
            if template_tr is not None:
                new_tr = copy.deepcopy(template_tr)
                table._tbl.append(new_tr)
                new_row = table.rows[-1]
            else:
                new_row = table.add_row()

            cells = new_row.cells
            n     = len(cells)

            # Columns 0–7: plain text, preserving template font
            plain_vals = [
                record['ntc_id'],   # 0 Notice ID
                record['adm'],      # 1 ADM
                record['sat_name'], # 2 STATION
                record['type'],     # 3 TYPE (G/N)
                record['d_rcv'],    # 4 Date of Receipt
                adm_coord,          # 5 ADM Coordination completed
                '',                 # 6 ADM Coordination no longer required
                findings_review,    # 7 Findings review required (13A)
            ]
            for i, val in enumerate(plain_vals):
                if i < n:
                    _set_cell_text(cells[i], val)

            # Column 8: 11.41 still exist + superscript footnote marker
            if 8 < n:
                _set_cell_still_exist(cells[8], mdb_data['still_exist'], still_exist_suffix)

            # Column 9: Provision updated from 11.41|X| – explicit Calibri
            if 9 < n:
                _set_cell_text(cells[9], provn_col, force_calibri=True)


def generate_report(
    notice_ids: list,
    wm_num: str,
    d_num: str,
    meeting_date: str,
) -> tuple:
    """
    Generate the Word report for one or more notice IDs.
    Returns (BytesIO buffer, filename string).
    """
    # ── Gather data for all notices ───────────────────────────────────────────
    notice_data_list = []
    all_provn_targets: set = set()

    for notice_id in notice_ids:
        mdb_data     = get_mdb_data(notice_id)
        history_data = get_history_data(notice_id)
        capture_data = get_capture_data(notice_id)
        notice_data_list.append((mdb_data, history_data, capture_data))
        all_provn_targets.update(history_data['provn_targets'])

    combined_provn_targets = sorted(all_provn_targets, key=_numeric_key)

    # ── Load a fresh copy of the template ────────────────────────────────────
    doc = Document(TEMPLATE_PATH)

    # ── 1. Update meeting header paragraph ───────────────────────────────────
    for para in doc.paragraphs:
        if 'meeting on' in para.text:
            _update_meeting_para(para, wm_num, meeting_date)
            break

    # ── 2. Replace provn_target section in coordination paragraph ────────────
    for para in doc.paragraphs:
        if 'coordination agreement with the concerned administrations' in para.text:
            _replace_provn_section(para, combined_provn_targets)
            break

    # ── 3. Replace table rows ─────────────────────────────────────────────────
    if doc.tables:
        _replace_table_data(doc.tables[0], notice_data_list)

    # ── 4. Update footer with document numbers ────────────────────────────────
    _update_footer(doc, wm_num, d_num)

    # ── 5. Update sign-off date ───────────────────────────────────────────────
    today_str = (
        datetime.today().strftime('%#d %B %Y') if os.name == 'nt'
        else datetime.today().strftime('%-d %B %Y')
    )
    _update_signoff_date(doc, today_str)

    # ── 6. Determine output filename ─────────────────────────────────────────
    wm_part = wm_num  if wm_num  else 'xxxx'
    d_part  = d_num   if d_num   else 'xxxxx'
    filename = f'11.41B_WM{wm_part}_D{d_part}_Draft.docx'

    # ── 7. Save to buffer and return ─────────────────────────────────────────
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf, filename


# ── Flask routes ──────────────────────────────────────────────────────────────

@app.route('/', methods=['GET'])
def index():
    # Enumerate available notice IDs from the base directory
    notice_ids = []
    if os.path.isdir(BASE_DIR):
        for name in sorted(os.listdir(BASE_DIR)):
            folder = os.path.join(BASE_DIR, name)
            mdb    = os.path.join(folder, f'{name}.mdb')
            if os.path.isdir(folder) and os.path.isfile(mdb):
                notice_ids.append(name)
    return render_template('index.html', notice_ids=notice_ids)


@app.route('/generate', methods=['POST'])
def generate():
    # Support both multi-select (getlist) and comma-separated text input
    raw_ids    = request.form.getlist('notice_id')
    notice_ids = []
    for raw in raw_ids:
        for nid in raw.split(','):
            nid = nid.strip()
            if nid:
                notice_ids.append(nid)

    wm_num       = (request.form.get('wm_num', '')    or '').strip()
    d_num        = (request.form.get('d_num', '')     or '').strip()
    meeting_date = (request.form.get('meeting_date', '') or '').strip()

    if not notice_ids:
        return jsonify({'error': 'At least one Notice ID is required.'}), 400

    # Default meeting date to today
    if not meeting_date:
        meeting_date = datetime.today().strftime('%-d %B %Y') if os.name != 'nt' \
            else datetime.today().strftime('%#d %B %Y')

    try:
        buf, filename = generate_report(notice_ids, wm_num, d_num, meeting_date)
    except FileNotFoundError as exc:
        return jsonify({'error': str(exc)}), 404
    except Exception as exc:
        return jsonify({'error': f'Unexpected error: {exc}'}), 500

    return send_file(
        buf,
        as_attachment=True,
        download_name=filename,
        mimetype=(
            'application/vnd.openxmlformats-officedocument'
            '.wordprocessingml.document'
        ),
    )


if __name__ == '__main__':
    app.run(debug=True, port=5050)
