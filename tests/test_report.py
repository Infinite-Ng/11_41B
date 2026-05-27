"""
Tests for 11.41B report-generator helper functions.

Covers the 4 runtime bugs that were fixed:
  Bug 1 – footer D-number update works even when the footer is a table
  Bug 2 – sign-off date replacement works when the date spans multiple runs
  Bug 3 – still-exist superscript is written correctly via _make_run_xml
  Bug 4 – meeting-heading search reaches paragraphs inside table cells
"""
import re
import sys
import os

import pytest

# ── make app.py importable without Flask complaining about templates ──────────
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from app import (
    _all_doc_paragraphs,
    _make_run_xml,
    _set_cell_still_exist,
    _update_footer,
    _update_meeting_para,
    _update_signoff_date,
    ordinal,
    ordinal_suffix,
)

W  = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
NS = 'http://www.w3.org/XML/1998/namespace'


# ─────────────────────────────────────────────────────────────────────────────
# Ordinal helper
# ─────────────────────────────────────────────────────────────────────────────

class TestOrdinal:
    @pytest.mark.parametrize('n, expected', [
        (1,    '1st'),
        (2,    '2nd'),
        (3,    '3rd'),
        (4,    '4th'),
        (11,   '11th'),
        (12,   '12th'),
        (13,   '13th'),
        (21,   '21st'),
        (22,   '22nd'),
        (101,  '101st'),
        (1668, '1668th'),
    ])
    def test_ordinal(self, n, expected):
        assert ordinal(n) == expected


# ─────────────────────────────────────────────────────────────────────────────
# Footer update (Bug 1)
# ─────────────────────────────────────────────────────────────────────────────

def _footer_all_text(doc: Document) -> str:
    """Collect all w:t text from the first section's default footer XML."""
    ftr = doc.sections[0].footer._element
    return ''.join((t.text or '') for t in ftr.iter(f'{{{W}}}t'))


def _make_doc_footer_paragraph(text: str) -> Document:
    """Document whose footer has one plain text paragraph."""
    doc = Document()
    section = doc.sections[0]
    section.footer.is_linked_to_previous = False
    section.footer.paragraphs[0].text = text
    return doc


def _make_doc_footer_table(text: str) -> Document:
    """Document whose footer holds a 1×1 TABLE (not a paragraph)."""
    doc = Document()
    section = doc.sections[0]
    section.footer.is_linked_to_previous = False
    ftr = section.footer._element
    # Build w:tbl > w:tr > w:tc > w:p > w:r > w:t manually
    tbl = OxmlElement('w:tbl')
    tr  = OxmlElement('w:tr')
    tc  = OxmlElement('w:tc')
    p   = OxmlElement('w:p')
    r   = OxmlElement('w:r')
    t   = OxmlElement('w:t')
    t.text = text
    r.append(t)
    p.append(r)
    tc.append(p)
    tr.append(tc)
    tbl.append(tr)
    ftr.append(tbl)
    return doc


class TestUpdateFooter:

    def test_plain_paragraph_updated(self):
        doc = _make_doc_footer_paragraph('WM1668 D72317')
        _update_footer(doc, '1668', '99999')
        assert 'D 99999' in _footer_all_text(doc)

    def test_table_footer_updated(self):
        """Bug 1: footer inside a table was never touched before the fix."""
        doc = _make_doc_footer_table('Report D72317 page 1')
        _update_footer(doc, '1668', '88888')
        text = _footer_all_text(doc)
        assert 'D 88888' in text
        assert 'D72317'  not in text

    def test_empty_d_num_leaves_footer_unchanged(self):
        doc = _make_doc_footer_paragraph('D 12345')
        _update_footer(doc, '1668', '')
        assert 'D 12345' in _footer_all_text(doc)

    def test_letter_prefix_not_mangled(self):
        """'ADMIN' or 'BD' must not be mis-matched by the regex."""
        doc = _make_doc_footer_paragraph('ADMIN D72317')
        _update_footer(doc, '1668', '00001')
        text = _footer_all_text(doc)
        assert 'ADMIN'   in text
        assert 'D 00001' in text


# ─────────────────────────────────────────────────────────────────────────────
# Sign-off date update (Bug 2)
# ─────────────────────────────────────────────────────────────────────────────

def _make_doc_date_para(run_texts: list) -> Document:
    """Document with one paragraph whose runs are the given strings."""
    doc  = Document()
    para = doc.add_paragraph()
    for text in run_texts:
        para.add_run(text)
    return doc


class TestUpdateSignoffDate:

    def test_single_run_replaced(self):
        doc = _make_doc_date_para(['Geneva, 15 January 2025'])
        _update_signoff_date(doc, '5 June 2026')
        assert '5 June 2026'     in doc.paragraphs[-1].text
        assert '15 January 2025' not in doc.paragraphs[-1].text

    def test_multi_run_merged_and_replaced(self):
        """Bug 2: date split across 3 runs was not replaced before the fix."""
        doc = _make_doc_date_para(['Geneva, 15 ', 'January', ' 2025'])
        _update_signoff_date(doc, '5 June 2026')
        full = ''.join(r.text for p in doc.paragraphs for r in p.runs)
        assert '5 June 2026' in full
        # Original year must be gone
        assert '2025' not in full

    def test_only_last_date_para_is_updated(self):
        doc = Document()
        doc.add_paragraph('Meeting on 10 March 2026')
        doc.add_paragraph('Signed 15 April 2026')
        _update_signoff_date(doc, '1 June 2026')
        assert '1 June 2026'   in doc.paragraphs[-1].text
        assert '10 March 2026' in doc.paragraphs[-2].text

    def test_no_date_para_is_noop(self):
        doc = Document()
        doc.add_paragraph('No date here')
        _update_signoff_date(doc, '1 June 2026')
        assert 'No date here' in doc.paragraphs[-1].text


# ─────────────────────────────────────────────────────────────────────────────
# Still-exist superscript (Bug 3)
# ─────────────────────────────────────────────────────────────────────────────

def _fresh_cell():
    """Return a single table cell from a new in-memory document."""
    doc   = Document()
    table = doc.add_table(rows=1, cols=1)
    return table.rows[0].cells[0]


class TestSetCellStillExist:

    def test_superscript_run_created(self):
        """Bug 3: _make_run_xml guarantees correct vertAlign=superscript XML."""
        cell = _fresh_cell()
        _set_cell_still_exist(cell, 'YES', 'A1')
        p_elem = cell.paragraphs[0]._p
        runs = p_elem.findall(f'{{{W}}}r')
        assert len(runs) == 2
        rPr2 = runs[1].find(f'{{{W}}}rPr')
        assert rPr2 is not None, 'superscript run has no rPr'
        va = rPr2.find(f'{{{W}}}vertAlign')
        assert va is not None, 'no vertAlign element'
        assert va.get(qn('w:val')) == 'superscript'
        t2 = runs[1].find(f'{{{W}}}t')
        assert t2 is not None and t2.text == 'A1'

    def test_no_superscript_when_suffix_empty(self):
        cell = _fresh_cell()
        _set_cell_still_exist(cell, 'NO', '')
        p_elem = cell.paragraphs[0]._p
        runs = p_elem.findall(f'{{{W}}}r')
        assert len(runs) == 1
        t1 = runs[0].find(f'{{{W}}}t')
        assert t1 is not None and t1.text == 'NO'

    def test_main_text_correct(self):
        cell = _fresh_cell()
        _set_cell_still_exist(cell, 'YES', 'B3')
        p_elem = cell.paragraphs[0]._p
        runs = p_elem.findall(f'{{{W}}}r')
        t1 = runs[0].find(f'{{{W}}}t')
        assert t1 is not None and t1.text == 'YES'

    @pytest.mark.parametrize('suffix', ['A1', 'A2', 'A3', 'B1', 'B2', 'B3'])
    def test_all_valid_suffixes(self, suffix):
        cell = _fresh_cell()
        _set_cell_still_exist(cell, 'YES', suffix)
        p_elem = cell.paragraphs[0]._p
        runs = p_elem.findall(f'{{{W}}}r')
        assert len(runs) == 2
        t2 = runs[1].find(f'{{{W}}}t')
        assert t2 is not None and t2.text == suffix


# ─────────────────────────────────────────────────────────────────────────────
# Meeting-heading paragraph search (Bug 4)
# ─────────────────────────────────────────────────────────────────────────────

MTG_TEXT = 'Presented at the 1668th meeting on 1 January 2026'
MTG_RE   = re.compile(r'meeting\s+on\b', re.IGNORECASE)


class TestAllDocParagraphs:

    def test_yields_body_paragraphs(self):
        doc = Document()
        doc.add_paragraph('Hello')
        texts = [p.text for p in _all_doc_paragraphs(doc)]
        assert 'Hello' in texts

    def test_yields_table_cell_paragraphs(self):
        """Bug 4: doc.paragraphs misses table-cell paragraphs; _all_doc_paragraphs does not."""
        doc   = Document()
        table = doc.add_table(rows=1, cols=1)
        table.rows[0].cells[0].paragraphs[0].add_run(MTG_TEXT)
        texts = [p.text for p in _all_doc_paragraphs(doc)]
        assert any('meeting on' in t for t in texts)


class TestUpdateMeetingParaBody:

    def test_date_updated_in_body_paragraph(self):
        doc  = Document()
        para = doc.add_paragraph()
        para.add_run('Presented at the ')
        para.add_run('1668')
        para.add_run('th meeting on ')
        para.add_run('1 January 2026')
        _update_meeting_para(para, '1668', '4 June 2026')
        assert '4 June 2026' in para.text

    def test_ordinal_superscript_run_present(self):
        doc  = Document()
        para = doc.add_paragraph()
        para.add_run(MTG_TEXT)
        _update_meeting_para(para, '1668', '4 June 2026')
        p_elem = para._p
        runs = p_elem.findall(f'{{{W}}}r')
        # Expect exactly 3 runs: prefix / ordinal-suffix / rest
        assert len(runs) == 3
        # Second run must be superscript
        rPr = runs[1].find(f'{{{W}}}rPr')
        assert rPr is not None
        va = rPr.find(f'{{{W}}}vertAlign')
        assert va is not None and va.get(qn('w:val')) == 'superscript'

    def test_date_updated_in_table_cell_via_search(self):
        """Simulate the generate_report search loop now using _all_doc_paragraphs."""
        doc   = Document()
        table = doc.add_table(rows=1, cols=1)
        cell  = table.rows[0].cells[0]
        cell.paragraphs[0].add_run(MTG_TEXT)

        for p in _all_doc_paragraphs(doc):
            if MTG_RE.search(p.text):
                _update_meeting_para(p, '1668', '4 June 2026')
                break

        assert '4 June 2026' in cell.paragraphs[0].text
        assert '1 January 2026' not in cell.paragraphs[0].text
